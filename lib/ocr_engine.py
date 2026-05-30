import io
from PIL import Image

try:
    from paddleocr import PaddleOCR
    _PADDLEOCR_AVAILABLE = True
except ImportError:
    _PADDLEOCR_AVAILABLE = False

_IMAGE_EXTENSIONS = {"jpg", "jpeg", "png", "heic", "bmp", "tiff", "tif", "webp"}
_SUPPORTED_EXTENSIONS = _IMAGE_EXTENSIONS | {"pdf", "docx", "doc", "xls", "xlsx", "csv", "txt"}


def is_image_file(filename: str) -> bool:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    return ext in _IMAGE_EXTENSIONS


def is_supported_file(filename: str) -> bool:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    return ext in _SUPPORTED_EXTENSIONS


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext in _IMAGE_EXTENSIONS:
        image = Image.open(io.BytesIO(file_bytes))
        return _run_ocr(image)

    if ext == "pdf":
        text = ""
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except ImportError:
            pass
        except Exception:
            pass

        if text.strip():
            return text

        try:
            from pdf2image import convert_from_bytes
            images = convert_from_bytes(file_bytes)
            for img in images:
                page_text = _run_ocr(img)
                if page_text:
                    text += page_text + "\n"
        except ImportError:
            pass
        except Exception:
            pass

        return text

    if ext in ("docx", "doc"):
        text = ""
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs:
                if para.text:
                    text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    row_text = " | ".join(cells)
                    if row_text.strip("| "):
                        text += row_text + "\n"
        except ImportError:
            pass
        except Exception:
            pass
        return text

    return ""


def extract_receipt_data(file_bytes: bytes, filename: str = None) -> dict:
    result = {
        "merchant": "",
        "date": None,
        "total": 0.0,
        "tax": 0.0,
        "type": "miscellaneous",
        "description": "",
        "payment_method": "N/A",
    }

    try:
        from lib.receipt_parser import (
            extract_total, extract_date, extract_tax,
            classify_receipt_type, extract_merchant,
            extract_description, extract_payment_method,
        )
        _use_receipt_parser = True
    except ImportError:
        _use_receipt_parser = False

    try:
        if filename and not is_image_file(filename):
            ocr_text = extract_text_from_file(file_bytes, filename)
        else:
            image = Image.open(io.BytesIO(file_bytes))
            ocr_text = _run_ocr(image)

        if _use_receipt_parser:
            lines = ocr_text.strip().split("\n")
            result["merchant"] = extract_merchant(lines)
            date_str = extract_date(ocr_text)
            if date_str:
                try:
                    from datetime import datetime as dt
                    result["date"] = dt.strptime(date_str, "%Y-%m-%d").date()
                except ValueError:
                    result["date"] = None
            else:
                result["date"] = None
            result["total"] = extract_total(ocr_text) or 0.0
            result["tax"] = extract_tax(ocr_text) or 0.0
            result["type"] = classify_receipt_type(ocr_text, result.get("merchant", ""))
            result["description"] = extract_description(lines)
            result["payment_method"] = extract_payment_method(ocr_text)
        else:
            result = _parse_ocr_result(ocr_text, result)
    except Exception:
        pass

    return result


def _run_ocr(image: Image.Image) -> str:
    text = ""

    if _PADDLEOCR_AVAILABLE:
        try:
            ocr = PaddleOCR(use_angle_cls=True, lang="ch")
            img_array = __import__("numpy").array(image)
            ocr_result = ocr.ocr(img_array, cls=True)
            for line in ocr_result:
                if line:
                    for word_info in line:
                        text += word_info[1][0] + "\n"
        except Exception:
            text = ""

    if not text.strip():
        try:
            import pytesseract
            text = pytesseract.image_to_string(image, lang="chi_sim+eng")
        except ImportError:
            pass
        except Exception:
            pass

    return text


def _parse_ocr_result(text: str, result: dict) -> dict:
    import re

    lines = text.strip().split("\n")
    if lines:
        result["merchant"] = lines[0].strip()

    date_patterns = [
        r"(\d{4}[/\-]\d{1,2}[/\-]\d{1,2})",
        r"(\d{1,2}[/\-]\d{1,2}[/\-]\d{4})",
        r"(\d{1,2}[/\-]\d{1,2}[/\-]\d{2})",
    ]
    for pattern in date_patterns:
        match = re.search(pattern, text)
        if match:
            from datetime import datetime
            try:
                date_str = match.group(1)
                for fmt in ["%Y/%m/%d", "%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%d/%m/%y"]:
                    try:
                        result["date"] = datetime.strptime(date_str, fmt).date()
                        break
                    except ValueError:
                        continue
            except Exception:
                pass
            break

    total_patterns = [
        r"(?:TOTAL|Total|總計|合計|總額|Amount)[\s:]*\$?\s*([\d,]+\.?\d*)",
        r"\$\s*([\d,]+\.?\d*)\s*$",
    ]
    for pattern in total_patterns:
        match = re.search(pattern, text, re.MULTILINE)
        if match:
            try:
                result["total"] = float(match.group(1).replace(",", ""))
            except ValueError:
                pass
            break

    tax_patterns = [
    ]
    for pattern in tax_patterns:
        match = re.search(pattern, text)
        if match:
            try:
                result["tax"] = float(match.group(1).replace(",", ""))
            except ValueError:
                pass
            break

    text_lower = text.lower()
    if any(kw in text_lower for kw in ["restaurant", "餐廳", "茶餐", "酒樓", "快餐"]):
        result["type"] = "meals_entertainment"
    elif any(kw in text_lower for kw in ["mtr", "巴士", "的士", "taxi", "transport"]):
        result["type"] = "transportation"
    elif any(kw in text_lower for kw in ["電力", "煤氣", "水費", "electricity", "gas", "water", "中電", "港燈"]):
        result["type"] = "utilities"
    elif any(kw in text_lower for kw in ["超市", "便利店", "supermarket", "7-eleven", "wellcome", "parknshop"]):
        result["type"] = "miscellaneous"
    elif any(kw in text_lower for kw in ["辦公", "文具", "stationery"]):
        result["type"] = "office_supplies"
    elif any(kw in text_lower for kw in ["租金", "差餉", "rent", "rates"]):
        result["type"] = "rent_rates"
    elif any(kw in text_lower for kw in ["專業", "professional", "法律", "legal", "會計", "accounting"]):
        result["type"] = "professional_fees"
    elif any(kw in text_lower for kw in ["保險", "insurance"]):
        result["type"] = "insurance"
    elif any(kw in text_lower for kw in ["維修", "repair", "保養", "maintenance"]):
        result["type"] = "repairs_maintenance"
    elif any(kw in text_lower for kw in ["機票", "flight", "酒店", "hotel"]):
        result["type"] = "travel"
    elif any(kw in text_lower for kw in ["廣告", "advertising", "推廣", "promotion"]):
        result["type"] = "marketing"
    elif any(kw in text_lower for kw in ["折舊", "depreciation"]):
        result["type"] = "depreciation"

    if len(lines) > 1:
        for line in lines[1:]:
            cleaned = line.strip()
            if cleaned and len(cleaned) >= 2:
                result["description"] = cleaned
                break

    payment_keywords = [
        (r'\bVisa\b', "Visa"), (r'\bMastercard\b', "Mastercard"),
        (r'八達通|\bOctopus\b', "Octopus"), (r'現金|\bCash\b', "Cash"),
        (r'信用卡|\bCredit\s*Card\b', "Credit Card"),
        (r'易辦事|\bEPS\b', "EPS"), (r'\bPayMe\b', "PayMe"),
        (r'支付寶|\bAlipay\b', "Alipay"), (r'微信支付|\bWeChat\s*Pay\b', "WeChat Pay"),
        (r'轉數快|\bFPS\b', "FPS"),
    ]
    for pattern, label in payment_keywords:
        if re.search(pattern, text, re.IGNORECASE):
            result["payment_method"] = label
            break

    return result
